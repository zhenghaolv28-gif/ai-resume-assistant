from __future__ import annotations

from importlib.util import module_from_spec, spec_from_file_location
from io import BytesIO
from pathlib import Path
import unittest
from unittest.mock import patch

from docx import Document
from reportlab.pdfgen.canvas import Canvas


MODULE_PATH = Path(__file__).resolve().parents[1] / "resume_import.py"
SPEC = spec_from_file_location("resume_import_under_test", MODULE_PATH)
assert SPEC and SPEC.loader
RESUME_IMPORT = module_from_spec(SPEC)
SPEC.loader.exec_module(RESUME_IMPORT)


class ResumeParserTests(unittest.TestCase):
    def test_text_pdf_does_not_start_ocr(self) -> None:
        native_text = (
            "张三\n求职意向：产品经理\n工作经历\n"
            "2022-至今 某科技公司 产品经理\n教育经历\n某大学 本科"
        )
        with patch.object(RESUME_IMPORT, "_read_pdf_native", return_value=native_text), patch.object(
            RESUME_IMPORT,
            "_read_pdf_ocr",
            side_effect=AssertionError("普通文字 PDF 不应调用 OCR"),
        ):
            result = RESUME_IMPORT.extract_resume_text_with_details(b"pdf", "resume.pdf")

        self.assertFalse(result.ocr_used)
        self.assertEqual(result.method, "PDF 可复制文字")
        self.assertEqual(result.text, native_text)

    def test_scanned_pdf_automatically_falls_back_to_ocr(self) -> None:
        ocr_text = (
            "李雷\n目标岗位：数据分析师\n工作经历\n"
            "2023-至今 某零售公司 数据分析师\n专业技能\nPython、SQL"
        )
        with patch.object(RESUME_IMPORT, "_read_pdf_native", return_value=""), patch.object(
            RESUME_IMPORT,
            "_read_pdf_ocr",
            return_value=(ocr_text, "chi_sim+eng"),
        ):
            result = RESUME_IMPORT.extract_resume_text_with_details(b"scan", "scan.pdf")

        self.assertTrue(result.ocr_used)
        self.assertEqual(result.method, "扫描 PDF OCR")
        self.assertEqual(result.ocr_language, "chi_sim+eng")
        self.assertIn("数据分析师", result.text)

    def test_image_resume_uses_ocr_and_existing_parser(self) -> None:
        ocr_text = (
            "王芳\n所在城市：深圳\n求职意向：用户运营\n"
            "工作经历\n2024-至今 某互联网公司 用户运营"
        )
        with patch.object(
            RESUME_IMPORT,
            "_read_image_ocr",
            return_value=(ocr_text, "chi_sim+eng"),
        ):
            result = RESUME_IMPORT.extract_resume_text_with_details(b"image", "resume.PNG")

        parsed = RESUME_IMPORT.parse_resume_text(result.text)
        self.assertTrue(result.ocr_used)
        self.assertEqual(result.method, "图片 OCR")
        self.assertEqual(parsed["city"], "深圳")
        self.assertIn("某互联网公司", parsed["work_experience"])

    def test_ocr_result_with_too_little_text_has_clear_error(self) -> None:
        with patch.object(
            RESUME_IMPORT,
            "_read_image_ocr",
            return_value=("空白", "chi_sim+eng"),
        ):
            with self.assertRaisesRegex(ValueError, "OCR 没有识别到足够文字"):
                RESUME_IMPORT.extract_resume_text_with_details(b"image", "blank.jpg")

    def test_plain_contact_header_is_extracted_and_not_duplicated(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """张三
13800138000 | zhangsan@example.com | 上海
求职意向：产品经理
个人优势
三年互联网产品经验。
工作经历
2022.03-至今 某科技公司 产品经理
教育背景
某大学 本科
专业技能
Axure | SQL | 数据分析"""
        )

        self.assertEqual(parsed["name"], "张三")
        self.assertEqual(parsed["phone"], "13800138000")
        self.assertEqual(parsed["email"], "zhangsan@example.com")
        self.assertEqual(parsed["city"], "上海")
        self.assertEqual(parsed["skills"], "Axure | SQL | 数据分析")
        self.assertEqual(parsed["unclassified_text"], "")
        self.assertEqual(parsed["parse_quality"], "高")

    def test_composite_labeled_header(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """姓名：李雷 | 手机：13900139000 | 邮箱：lilei@example.com | 现居地：杭州
目标岗位：数据分析师
职业摘要：熟悉 SQL、Python 和业务分析。
工作经验
2021.07-至今 某零售公司 数据分析师
教育经历
某大学 统计学 本科"""
        )

        self.assertEqual(parsed["name"], "李雷")
        self.assertEqual(parsed["city"], "杭州")
        self.assertEqual(parsed["target_role"], "数据分析师")
        self.assertEqual(parsed["unclassified_text"], "")

    def test_space_separated_labels(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """姓名 王芳
联系电话 13700137000
电子邮箱 wangfang@example.com
所在城市 深圳
期望岗位 用户运营
自我评价
有用户增长和活动运营经验。
任职经历
2020.06-至今 某互联网公司 用户运营"""
        )

        self.assertEqual(parsed["name"], "王芳")
        self.assertEqual(parsed["city"], "深圳")
        self.assertEqual(parsed["target_role"], "用户运营")
        self.assertEqual(parsed["summary"], "有用户增长和活动运营经验。")

    def test_city_inside_compact_personal_header(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """王芳
女 / 25岁 / 四川成都
求职意向：用户运营
工作经历
2023-至今 某互联网公司 用户运营"""
        )

        self.assertEqual(parsed["city"], "成都")

    def test_city_from_full_prefecture_list_without_city_suffix(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """王芳
女 / 25岁 / 安徽 蚌埠
求职意向：用户运营
工作经历
2023-至今 某互联网公司 用户运营"""
        )

        self.assertEqual(parsed["city"], "蚌埠")

    def test_city_from_province_city_district_label(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """姓名：李雷
现居住地广东省深圳市南山区
目标岗位：数据分析师
工作经历
2022-至今 某科技公司 数据分析师"""
        )

        self.assertEqual(parsed["city"], "深圳")

    def test_less_common_city_is_kept_from_label(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """赵敏
常住地：邯郸
期望岗位：产品经理
工作经历
2022-至今 某公司 产品专员"""
        )

        self.assertEqual(parsed["city"], "邯郸")

    def test_english_labeled_city(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """Alex Chen
Location Shanghai, China
Target Position: Product Manager
Professional Experience
2022-Present Example Labs"""
        )

        self.assertEqual(parsed["city"], "Shanghai")

    def test_interleaved_summary_content_is_removed_from_skills(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """王芳
自我介绍
专业技能
本人拥有三年用户运营经验，责任心强，善于沟通与团队协作。
Excel、PowerPoint、SQL、数据分析
工作经历
2022-至今 某互联网公司 用户运营"""
        )

        self.assertEqual(
            parsed["summary"],
            "本人拥有三年用户运营经验,责任心强,善于沟通与团队协作。",
        )
        self.assertEqual(parsed["skills"], "Excel、PowerPoint、SQL、数据分析")

    def test_skill_prose_is_not_moved_to_summary(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """李雷
专业技能
熟练掌握 Python、SQL，具有3年项目使用经验
熟练使用 Excel 与 PowerPoint
工作经历
2022-至今 某科技公司 数据分析师"""
        )

        self.assertEqual(parsed["summary"], "")
        self.assertIn("熟练掌握 Python", parsed["skills"])

    def test_english_resume_and_international_phone(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """Alex Chen | Singapore | +65 8123 4567 | alex.chen@example.com
TARGET POSITION: Product Manager
PROFESSIONAL SUMMARY
Product manager with five years of B2B SaaS experience.
PROFESSIONAL EXPERIENCE
2021-Present Example Labs, Product Manager
EDUCATION
National University of Singapore, B.Comp.
TECHNICAL SKILLS
SQL, Figma, Product Analytics"""
        )

        self.assertEqual(parsed["name"], "Alex Chen")
        self.assertEqual(parsed["phone"], "+65 8123 4567")
        self.assertEqual(parsed["city"], "Singapore")
        self.assertEqual(parsed["target_role"], "Product Manager")
        self.assertEqual(parsed["parse_quality"], "高")

    def test_campus_experience_is_kept_with_experience(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """赵敏
手机：13600136000
求职目标：前端工程师
工作经历
2023.01-至今 某软件公司
校园经历
学生会技术部负责人
教育经历
某大学 软件工程 本科"""
        )

        self.assertIn("学生会技术部负责人", parsed["work_experience"])
        self.assertNotIn("校园经历", parsed["unclassified_text"])

    def test_low_confidence_result_has_actionable_warnings(self) -> None:
        parsed = RESUME_IMPORT.parse_resume_text(
            """候选人资料
这是一段没有标准章节、没有联系电话、也没有电子邮箱的长篇说明文字。
程序应该把这部分保留为未分类内容，不能假装已经准确识别。"""
        )

        self.assertEqual(parsed["parse_quality"], "低")
        self.assertGreaterEqual(len(parsed["parse_warnings"]), 2)
        self.assertTrue(parsed["unclassified_text"])

    def test_docx_paragraph_and_table_content_is_read(self) -> None:
        document = Document()
        document.add_paragraph("张三")
        document.add_paragraph("工作经历")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "2022-至今"
        table.cell(0, 1).text = "某科技公司 产品经理"
        document.add_paragraph("教育经历")
        document.add_paragraph("某大学 本科")
        buffer = BytesIO()
        document.save(buffer)

        extracted = RESUME_IMPORT.extract_resume_text(buffer.getvalue(), "resume.docx")
        parsed = RESUME_IMPORT.parse_resume_text(extracted)

        self.assertIn("某科技公司", parsed["work_experience"])
        self.assertIn("某大学", parsed["education"])

    def test_two_column_docx_keeps_every_section_complete(self) -> None:
        document = Document()
        document.add_paragraph("Alex Chen")
        table = document.add_table(rows=4, cols=2)
        table.cell(0, 0).text = "INTERNSHIP EXPERIENCE"
        table.cell(0, 1).text = "TECHNICAL SKILLS"
        table.cell(1, 0).text = "2024 Example Labs, Product Intern"
        table.cell(1, 1).text = "Python, SQL, Figma"
        table.cell(2, 0).text = "EDUCATION"
        table.cell(2, 1).text = "PROJECT EXPERIENCE"
        table.cell(3, 0).text = "Example University, B.Sc."
        table.cell(3, 1).text = "Resume parser project"
        buffer = BytesIO()
        document.save(buffer)

        extracted = RESUME_IMPORT.extract_resume_text(buffer.getvalue(), "two-column.docx")
        parsed = RESUME_IMPORT.parse_resume_text(extracted)

        self.assertEqual(parsed["work_experience"], "2024 Example Labs, Product Intern")
        self.assertEqual(parsed["skills"], "Python, SQL, Figma")
        self.assertEqual(parsed["education"], "Example University, B.Sc.")
        self.assertEqual(parsed["project_experience"], "Resume parser project")

    def test_chinese_two_column_docx_separates_internship_and_skills(self) -> None:
        document = Document()
        document.add_paragraph("王芳")
        table = document.add_table(rows=4, cols=2)
        table.cell(0, 0).text = "实习经历"
        table.cell(0, 1).text = "专业技能"
        table.cell(1, 0).text = "2024.06-2024.12 某科技公司 产品实习生"
        table.cell(1, 1).text = "Python、SQL、数据分析"
        table.cell(2, 0).text = "教育经历"
        table.cell(2, 1).text = "项目经历"
        table.cell(3, 0).text = "某大学 信息管理 本科"
        table.cell(3, 1).text = "智能简历助手项目"
        buffer = BytesIO()
        document.save(buffer)

        extracted = RESUME_IMPORT.extract_resume_text(buffer.getvalue(), "中文两栏简历.docx")
        parsed = RESUME_IMPORT.parse_resume_text(extracted)

        self.assertEqual(
            parsed["work_experience"],
            "2024.06-2024.12 某科技公司 产品实习生",
        )
        self.assertEqual(parsed["skills"], "Python、SQL、数据分析")
        self.assertEqual(parsed["education"], "某大学 信息管理 本科")
        self.assertEqual(parsed["project_experience"], "智能简历助手项目")

    def test_two_column_pdf_keeps_every_section_complete(self) -> None:
        buffer = BytesIO()
        canvas = Canvas(buffer, pagesize=(612, 792))
        canvas.drawString(50, 750, "Alex Chen")
        left = [
            "INTERNSHIP EXPERIENCE",
            "2024 Example Labs, Product Intern",
            "EDUCATION",
            "Example University, B.Sc.",
        ]
        right = [
            "TECHNICAL SKILLS",
            "Python, SQL, Figma",
            "PROJECT EXPERIENCE",
            "Resume parser project",
        ]
        for index, (left_text, right_text) in enumerate(zip(left, right)):
            y = 700 - index * 28
            canvas.drawString(50, y, left_text)
            canvas.drawString(330, y, right_text)
        canvas.save()

        extracted = RESUME_IMPORT.extract_resume_text(buffer.getvalue(), "two-column.pdf")
        parsed = RESUME_IMPORT.parse_resume_text(extracted)

        self.assertEqual(parsed["work_experience"], "2024 Example Labs, Product Intern")
        self.assertEqual(parsed["skills"], "Python, SQL, Figma")
        self.assertEqual(parsed["education"], "Example University, B.Sc.")
        self.assertEqual(parsed["project_experience"], "Resume parser project")

    def test_normal_work_table_stays_in_row_order(self) -> None:
        document = Document()
        document.add_paragraph("张三")
        document.add_paragraph("工作经历")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "2022-2023"
        table.cell(0, 1).text = "甲公司 产品实习生"
        table.cell(1, 0).text = "2023-至今"
        table.cell(1, 1).text = "乙公司 产品经理"
        document.add_paragraph("专业技能")
        document.add_paragraph("Axure、SQL")
        buffer = BytesIO()
        document.save(buffer)

        extracted = RESUME_IMPORT.extract_resume_text(buffer.getvalue(), "work-table.docx")
        parsed = RESUME_IMPORT.parse_resume_text(extracted)

        self.assertEqual(
            parsed["work_experience"],
            "2022-2023\n甲公司 产品实习生\n2023-至今\n乙公司 产品经理",
        )
        self.assertEqual(parsed["skills"], "Axure、SQL")

    def test_repeated_long_duty_text_is_not_removed(self) -> None:
        repeated_duty = "负责用户需求分析、产品方案设计以及跨部门项目推进工作"
        document = Document()
        document.add_paragraph("张三")
        document.add_paragraph("工作经历")
        document.add_paragraph("2022-2023 甲公司 产品专员")
        document.add_paragraph(repeated_duty)
        document.add_paragraph("2023-至今 乙公司 产品经理")
        document.add_paragraph(repeated_duty)
        buffer = BytesIO()
        document.save(buffer)

        extracted = RESUME_IMPORT.extract_resume_text(buffer.getvalue(), "重复职责.docx")
        parsed = RESUME_IMPORT.parse_resume_text(extracted)

        self.assertEqual(parsed["work_experience"].count(repeated_duty), 2)

    def test_duplicate_document_block_is_collapsed_once(self) -> None:
        block = [
            "张三",
            "工作经历",
            "2022-至今 某科技公司 产品经理",
            "教育经历",
            "某大学 本科",
            "专业技能",
            "Axure、SQL、数据分析",
        ]
        document = Document()
        for line in block + block:
            document.add_paragraph(line)
        buffer = BytesIO()
        document.save(buffer)

        extracted = RESUME_IMPORT.extract_resume_text(buffer.getvalue(), "重复全文.docx")
        parsed = RESUME_IMPORT.parse_resume_text(extracted)

        self.assertEqual(parsed["work_experience"].count("某科技公司"), 1)
        self.assertEqual(parsed["education"].count("某大学"), 1)
        self.assertEqual(parsed["skills"].count("Axure"), 1)

    def test_duplicate_pdf_style_text_block_is_collapsed_once(self) -> None:
        block = [
            "李雷",
            "工作经历",
            "2021-至今 某零售公司 数据分析师",
            "教育经历",
            "某大学 统计学 本科",
            "专业技能",
            "Python、SQL",
        ]

        parsed = RESUME_IMPORT.parse_resume_text("\n".join(block + block))

        self.assertEqual(parsed["work_experience"].count("某零售公司"), 1)
        self.assertEqual(parsed["education"].count("某大学"), 1)
        self.assertEqual(parsed["skills"].count("Python"), 1)


if __name__ == "__main__":
    unittest.main()
