from io import BytesIO

from docx import Document
from pypdf import PdfReader

from resume_template import (
    DEFAULT_TEMPLATE_ID,
    RESUME_TEMPLATES,
    create_resume_document,
    create_resume_pdf,
    create_resume_preview_html,
    resume_template_options,
)


SAMPLE_RESUME = {
    "name": "林晓然",
    "target_role": "产品经理",
    "phone": "13800138000",
    "email": "lin@example.com",
    "city": "上海",
    "summary": "具备五年企业服务产品经验，擅长需求分析与跨团队协作。",
    "education": "2015.09-2019.06 同济大学 信息管理 本科",
    "work_experience": "2021.03-至今 某科技公司 产品经理\n- 负责产品规划与版本推进",
    "project_experience": "智能招聘平台｜项目负责人\n- 完成核心流程设计",
    "skills": "- Axure、Figma、SQL",
}


def test_four_professional_templates_are_available():
    options = resume_template_options()

    assert list(options) == [
        "business_blue",
        "minimal_mono",
        "executive_navy",
        "modern_teal",
    ]
    assert DEFAULT_TEMPLATE_ID in options


def test_preview_uses_selected_template_and_escapes_resume_content():
    resume = SAMPLE_RESUME | {"city": "上海 & 苏州"}

    preview = create_resume_preview_html(resume, template_id="modern_teal")

    assert RESUME_TEMPLATES["modern_teal"]["primary"] in preview
    assert "青绿现代简历实时预览" in preview
    assert "林晓然" in preview
    assert "上海 &amp; 苏州" in preview
    assert "上海 & 苏州" not in preview


def test_every_template_generates_readable_pdf_and_word():
    generated_pdfs = []
    for template_id, template_name in resume_template_options().items():
        pdf_bytes = create_resume_pdf(SAMPLE_RESUME, template_id=template_id)
        generated_pdfs.append(pdf_bytes)
        reader = PdfReader(BytesIO(pdf_bytes))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "简历" in text
        assert "林晓然" in text
        assert "工作经历" in text
        assert template_name in text

        word_bytes = create_resume_document(SAMPLE_RESUME, template_id=template_id)
        document = Document(BytesIO(word_bytes))
        document_text = "\n".join(
            paragraph.text for paragraph in document.paragraphs
        ) + "\n" + "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        assert "林晓然" in document_text
        assert "目标岗位：产品经理" in document_text

    assert len({pdf_bytes for pdf_bytes in generated_pdfs}) == 4


def test_unknown_template_falls_back_to_default():
    unknown_preview = create_resume_preview_html(SAMPLE_RESUME, template_id="missing")
    default_preview = create_resume_preview_html(
        SAMPLE_RESUME,
        template_id=DEFAULT_TEMPLATE_ID,
    )

    assert unknown_preview == default_preview
