"""生成带证件照的 Word 和 PDF 简历。"""

from __future__ import annotations

from io import BytesIO
import html
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    Image as PdfImage,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


RAW_SECTIONS = (
    ("职业概述", "summary"),
    ("教育经历", "education"),
    ("工作或实习经历", "work_experience"),
    ("项目经历", "project_experience"),
    ("技能与证书", "skills"),
)
CHINESE_FONT = "STSong-Light"


def _content_blocks(resume_data: dict, optimized_text: str | None) -> list[tuple[str, str]]:
    """把 AI Markdown 或原始资料整理成简单的标题、正文和项目符号。"""
    blocks: list[tuple[str, str]] = []
    if optimized_text and optimized_text.strip():
        for original_line in optimized_text.strip().splitlines():
            line = original_line.strip()
            if not line:
                continue
            if line.startswith("#"):
                title = re.sub(r"^#+\s*", "", line).strip()
                if title:
                    blocks.append(("heading", title))
            elif re.match(r"^[-*•]\s+", line):
                blocks.append(("bullet", re.sub(r"^[-*•]\s+", "", line)))
            else:
                blocks.append(("paragraph", line))
        return blocks

    for title, key in RAW_SECTIONS:
        content = str(resume_data.get(key, "")).strip()
        if not content:
            continue
        blocks.append(("heading", title))
        for line in content.splitlines():
            line = line.strip()
            if not line:
                continue
            if re.match(r"^[-*•]\s+", line):
                blocks.append(("bullet", re.sub(r"^[-*•]\s+", "", line)))
            else:
                blocks.append(("paragraph", line))
    return blocks


def _contact_line(resume_data: dict) -> str:
    return " ｜ ".join(
        str(item).strip()
        for item in (
            resume_data.get("phone", ""),
            resume_data.get("email", ""),
            resume_data.get("city", ""),
        )
        if str(item).strip()
    )


def _set_word_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold


def _add_word_text(paragraph, text: str, size: float = 10.5, bold: bool = False) -> None:
    _set_word_font(paragraph.add_run(text), size=size, bold=bold)


def create_resume_document(
    resume_data: dict,
    optimized_text: str | None = None,
    photo_bytes: bytes | None = None,
) -> bytes:
    """生成可下载的 Word 简历。"""
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(13.5)
    table.columns[1].width = Cm(3.2)
    left_cell, right_cell = table.rows[0].cells
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    name_paragraph = left_cell.paragraphs[0]
    _add_word_text(name_paragraph, str(resume_data.get("name", "简历")), 22, True)
    role_paragraph = left_cell.add_paragraph()
    _add_word_text(role_paragraph, f"目标岗位：{resume_data.get('target_role', '')}", 11, True)
    contact = _contact_line(resume_data)
    if contact:
        contact_paragraph = left_cell.add_paragraph()
        _add_word_text(contact_paragraph, contact, 9.5)

    if photo_bytes:
        photo_paragraph = right_cell.paragraphs[0]
        photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        photo_paragraph.add_run().add_picture(BytesIO(photo_bytes), width=Cm(2.8))

    for block_type, text in _content_blocks(resume_data, optimized_text):
        if block_type == "heading":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(4)
            _add_word_text(paragraph, text, 13, True)
        elif block_type == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            _add_word_text(paragraph, text)
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            _add_word_text(paragraph, text)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _register_pdf_font() -> None:
    try:
        pdfmetrics.getFont(CHINESE_FONT)
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont(CHINESE_FONT))


def _pdf_photo(photo_bytes: bytes) -> PdfImage:
    photo_buffer = BytesIO(photo_bytes)
    image_reader = ImageReader(photo_buffer)
    pixel_width, pixel_height = image_reader.getSize()
    max_width, max_height = 28 * mm, 38 * mm
    scale = min(max_width / pixel_width, max_height / pixel_height)
    photo_buffer.seek(0)
    return PdfImage(photo_buffer, width=pixel_width * scale, height=pixel_height * scale)


def create_resume_pdf(
    resume_data: dict,
    optimized_text: str | None = None,
    photo_bytes: bytes | None = None,
) -> bytes:
    """生成可直接下载的中文 PDF 简历。"""
    _register_pdf_font()
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=15 * mm,
        bottomMargin=16 * mm,
        title=f"{resume_data.get('name', '')}的简历",
        author=str(resume_data.get("name", "")),
    )

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "ResumeName",
        parent=styles["Title"],
        fontName=CHINESE_FONT,
        fontSize=22,
        leading=28,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#172B4D"),
        spaceAfter=6,
    )
    role_style = ParagraphStyle(
        "ResumeRole",
        parent=styles["Normal"],
        fontName=CHINESE_FONT,
        fontSize=11,
        leading=16,
        textColor=colors.HexColor("#334155"),
    )
    heading_style = ParagraphStyle(
        "ResumeHeading",
        parent=styles["Heading2"],
        fontName=CHINESE_FONT,
        fontSize=13,
        leading=18,
        textColor=colors.HexColor("#0F4C81"),
        spaceBefore=9,
        spaceAfter=4,
        borderWidth=0,
        borderPadding=0,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["BodyText"],
        fontName=CHINESE_FONT,
        fontSize=10.5,
        leading=16,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#202938"),
        spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet",
        parent=body_style,
        leftIndent=12,
        firstLineIndent=-8,
        bulletIndent=2,
    )

    header_items = [
        Paragraph(html.escape(str(resume_data.get("name", "简历"))), name_style),
        Paragraph(
            html.escape(f"目标岗位：{resume_data.get('target_role', '')}"),
            role_style,
        ),
    ]
    contact = _contact_line(resume_data)
    if contact:
        header_items.append(Paragraph(html.escape(contact), role_style))

    if photo_bytes:
        header = Table(
            [[header_items, _pdf_photo(photo_bytes)]],
            colWidths=[document.width - 34 * mm, 34 * mm],
        )
        header.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ]
            )
        )
        story = [header, Spacer(1, 4 * mm)]
    else:
        story = header_items + [Spacer(1, 4 * mm)]

    for block_type, text in _content_blocks(resume_data, optimized_text):
        safe_text = html.escape(text)
        if block_type == "heading":
            story.append(KeepTogether([Paragraph(safe_text, heading_style)]))
        elif block_type == "bullet":
            story.append(Paragraph(safe_text, bullet_style, bulletText="•"))
        else:
            story.append(Paragraph(safe_text, body_style))

    def add_page_number(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont(CHINESE_FONT, 8)
        canvas.setFillColor(colors.HexColor("#64748B"))
        canvas.drawCentredString(A4[0] / 2, 8 * mm, f"第 {doc.page} 页")
        canvas.restoreState()

    document.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return output.getvalue()
