"""招聘网站风格的中文简历模板，以及 Word/PDF 导出功能。"""

from __future__ import annotations

from base64 import b64encode
from io import BytesIO
import html
from pathlib import Path
import re
import unicodedata

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Flowable,
    Image as PdfImage,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


CHINESE_FONT = "STSong-Light"
EMBEDDED_CHINESE_FONT = "ResumeChinese"
EMBEDDED_CHINESE_BOLD_FONT = "ResumeChineseBold"
FONT_CANDIDATES = (
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
    Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
)
BOLD_FONT_CANDIDATES = (
    Path(r"C:\Windows\Fonts\msyhbd.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
    Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"),
)
DEFAULT_TEMPLATE_ID = "business_blue"
RESUME_TEMPLATES = {
    "business_blue": {
        "name": "蓝白商务",
        "description": "稳重清晰，适合大多数企业与通用岗位",
        "primary": "#123E67",
        "accent": "#2E6FA6",
        "soft": "#EAF3FB",
        "pale": "#F5F9FD",
        "border": "#C7DBEB",
        "text": "#1F2A37",
        "muted": "#5B6B7E",
        "header_mode": "light",
    },
    "minimal_mono": {
        "name": "黑白极简",
        "description": "留白克制，适合法律、咨询、研究与传统行业",
        "primary": "#18181B",
        "accent": "#3F3F46",
        "soft": "#F4F4F5",
        "pale": "#FAFAFA",
        "border": "#D4D4D8",
        "text": "#18181B",
        "muted": "#52525B",
        "header_mode": "line",
    },
    "executive_navy": {
        "name": "深蓝高管",
        "description": "沉稳有分量，适合管理、金融与资深专业岗位",
        "primary": "#102A43",
        "accent": "#D4A72C",
        "soft": "#E8EEF4",
        "pale": "#F3F6F9",
        "border": "#B9C7D5",
        "text": "#172B3A",
        "muted": "#526779",
        "header_mode": "dark",
    },
    "modern_teal": {
        "name": "青绿现代",
        "description": "清新理性，适合产品、设计、技术与创新团队",
        "primary": "#155E63",
        "accent": "#0F8B8D",
        "soft": "#DDF3F1",
        "pale": "#F1FAF9",
        "border": "#B9DEDB",
        "text": "#183536",
        "muted": "#526A6B",
        "header_mode": "light",
    },
}
SECTION_LABELS = {
    "职业概述": "个人简介",
    "个人简介": "个人简介",
    "自我介绍": "个人简介",
    "教育经历": "教育背景",
    "教育背景": "教育背景",
    "工作或实习经历": "工作经历",
    "工作经历": "工作经历",
    "项目经历": "项目经历",
    "技能与证书": "技能证书",
    "技能证书": "技能证书",
}
RAW_SECTIONS = (
    ("个人简介", "summary"),
    ("教育背景", "education"),
    ("工作经历", "work_experience"),
    ("项目经历", "project_experience"),
    ("技能证书", "skills"),
)
ALLOWED_SYMBOLS = set("+-@./\\%&")
MARKDOWN_SYMBOLS = set("*_`~#")


def _template(template_id: str | None) -> dict:
    """返回可用模板，未知值回退到默认模板。"""
    return RESUME_TEMPLATES.get(str(template_id or ""), RESUME_TEMPLATES[DEFAULT_TEMPLATE_ID])


def resume_template_options() -> dict[str, str]:
    """返回模板标识与界面显示名称。"""
    return {template_id: config["name"] for template_id, config in RESUME_TEMPLATES.items()}


def resume_template_description(template_id: str | None) -> str:
    return _template(template_id)["description"]


def _clean_line(text: str, preserve_bullet: bool = True) -> str:
    """清理 Markdown、表情和不可见字符，保留中文及常用简历标点。"""
    line = str(text).replace("\uFEFF", "").replace("\uFFFD", "")
    line = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", line)
    line = re.sub(r"^\s*#{1,6}\s*", "", line)
    line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
    line = re.sub(r"__(.*?)__", r"\1", line)
    is_bullet = bool(re.match(r"^\s*[-+*•·]\s+", line))
    line = re.sub(r"^\s*[-+*•·]\s+", "", line)

    cleaned: list[str] = []
    for character in line:
        if character in MARKDOWN_SYMBOLS:
            continue
        if character in "\r\n\t":
            cleaned.append(" ")
            continue
        category = unicodedata.category(character)
        if category.startswith("C") or category == "So":
            continue
        if character == "•":
            character = "·"
        if character.isalnum() or category.startswith("L") or category.startswith("N"):
            cleaned.append(character)
        elif category.startswith("P") or character in ALLOWED_SYMBOLS or character.isspace():
            cleaned.append(character)

    result = re.sub(r"[ \t]+", " ", "".join(cleaned)).strip()
    if preserve_bullet and is_bullet and result:
        result = f"· {result}"
    return result


def clean_resume_text(text: str | None) -> str:
    """清理 AI 结果，确保编辑框和导出文件不出现 Markdown 星号或乱码。"""
    if not text:
        return ""
    normalized = str(text).replace("\r\n", "\n").replace("\r", "\n")
    return "\n".join(_clean_line(line) for line in normalized.split("\n")).strip()


def _section_label(text: str) -> str:
    cleaned = _clean_line(text, preserve_bullet=False).rstrip("：:")
    return SECTION_LABELS.get(cleaned, cleaned)


def _content_blocks(resume_data: dict, optimized_text: str | None) -> list[tuple[str, str]]:
    """将 AI 文本或原始资料整理为招聘简历常见的标题、正文和项目符号。"""
    blocks: list[tuple[str, str]] = []
    if optimized_text and optimized_text.strip():
        for original_line in str(optimized_text).splitlines():
            raw_line = original_line.strip()
            if not raw_line:
                continue
            is_heading = bool(re.match(r"^\s*#{1,6}\s*", raw_line))
            is_bullet = bool(re.match(r"^\s*[-+*•·]\s+", raw_line))
            line = _clean_line(raw_line, preserve_bullet=False)
            if not line:
                continue
            if is_heading or line in SECTION_LABELS:
                blocks.append(("heading", _section_label(line)))
            elif is_bullet:
                blocks.append(("bullet", line))
            else:
                blocks.append(("paragraph", line))
        return blocks

    for title, key in RAW_SECTIONS:
        content = str(resume_data.get(key, "")).strip()
        if not content:
            continue
        blocks.append(("heading", title))
        for original_line in content.splitlines():
            raw_line = original_line.strip()
            if not raw_line:
                continue
            is_bullet = bool(re.match(r"^\s*[-+*•·]\s+", raw_line))
            line = _clean_line(raw_line, preserve_bullet=False)
            if line:
                blocks.append(("bullet" if is_bullet else "paragraph", line))
    return blocks


def _plain(value: object) -> str:
    return clean_resume_text(str(value or "")).replace("\n", " ").strip()


def _contact_line(resume_data: dict) -> str:
    items = [
        _plain(resume_data.get("phone")),
        _plain(resume_data.get("email")),
        _plain(resume_data.get("city")),
    ]
    return " ｜ ".join(item for item in items if item)


def create_resume_preview_html(
    resume_data: dict,
    optimized_text: str | None = None,
    photo_bytes: bytes | None = None,
    template_id: str = DEFAULT_TEMPLATE_ID,
) -> str:
    """生成与导出模板共享内容和配色的实时 HTML 预览。"""
    template = _template(template_id)
    name = html.escape(_plain(resume_data.get("name")) or "姓名")
    role = html.escape(_plain(resume_data.get("target_role")) or "目标岗位")
    contact = html.escape(_contact_line(resume_data))
    photo_html = ""
    if photo_bytes:
        encoded_photo = b64encode(photo_bytes).decode("ascii")
        photo_html = (
            '<img class="resume-preview-photo" '
            f'src="data:image/jpeg;base64,{encoded_photo}" alt="简历照片">'
        )

    body_parts: list[str] = []
    for block_type, text in _content_blocks(resume_data, optimized_text):
        safe_text = html.escape(text)
        if block_type == "heading":
            body_parts.append(f'<h3 class="resume-preview-section">{safe_text}</h3>')
        elif block_type == "bullet":
            body_parts.append(f'<p class="resume-preview-bullet">{safe_text}</p>')
        else:
            body_parts.append(f'<p class="resume-preview-paragraph">{safe_text}</p>')

    header_mode = template["header_mode"]
    header_class = f"resume-preview-header resume-preview-header-{header_mode}"
    return f"""
<div class="resume-preview-shell" style="
    --resume-primary:{template['primary']};
    --resume-accent:{template['accent']};
    --resume-soft:{template['soft']};
    --resume-pale:{template['pale']};
    --resume-border:{template['border']};
    --resume-text:{template['text']};
    --resume-muted:{template['muted']};">
  <style>
    .resume-preview-shell {{ width: 100%; padding: 18px; border-radius: 18px; background: #111318; box-sizing: border-box; }}
    .resume-preview-page {{ width: min(100%, 760px); min-height: 980px; margin: 0 auto; padding: 46px 52px 54px; box-sizing: border-box; background: #fff; color: var(--resume-text); box-shadow: 0 24px 60px rgba(0,0,0,.28); font-family: "Microsoft YaHei", "Noto Sans CJK SC", sans-serif; }}
    .resume-preview-title {{ margin: -46px -52px 24px; padding: 13px; text-align: center; color: #fff; background: var(--resume-primary); font-size: 20px; font-weight: 800; letter-spacing: .42em; text-indent: .42em; }}
    .resume-preview-header {{ display: flex; align-items: center; justify-content: space-between; gap: 24px; padding: 24px 26px; border: 1px solid var(--resume-border); border-bottom: 3px solid var(--resume-accent); background: var(--resume-pale); }}
    .resume-preview-header-dark {{ color: #fff; background: var(--resume-primary); border-color: var(--resume-primary); border-bottom-color: var(--resume-accent); }}
    .resume-preview-header-line {{ padding-left: 0; padding-right: 0; background: #fff; border-width: 0 0 2px; }}
    .resume-preview-name {{ margin: 0 0 7px; font-size: 31px; line-height: 1.18; font-weight: 800; letter-spacing: .04em; }}
    .resume-preview-role {{ margin: 0 0 7px; color: var(--resume-accent); font-size: 15px; font-weight: 700; }}
    .resume-preview-header-dark .resume-preview-role {{ color: var(--resume-soft); }}
    .resume-preview-contact {{ margin: 0; color: var(--resume-muted); font-size: 12px; line-height: 1.65; overflow-wrap: anywhere; }}
    .resume-preview-header-dark .resume-preview-contact {{ color: #dce7ef; }}
    .resume-preview-photo {{ width: 92px; height: 116px; object-fit: cover; border: 4px solid #fff; box-shadow: 0 3px 14px rgba(16,42,67,.16); flex: 0 0 auto; }}
    .resume-preview-body {{ padding-top: 13px; }}
    .resume-preview-section {{ margin: 20px 0 10px; padding: 9px 12px; color: var(--resume-primary); background: var(--resume-pale); border: 1px solid var(--resume-border); border-left: 5px solid var(--resume-accent); font-size: 17px; line-height: 1.35; }}
    .resume-preview-paragraph, .resume-preview-bullet {{ margin: 0 0 7px; font-size: 13.5px; line-height: 1.72; white-space: normal; overflow-wrap: anywhere; }}
    .resume-preview-bullet {{ position: relative; padding-left: 17px; }}
    .resume-preview-bullet::before {{ content: ""; position: absolute; left: 3px; top: .72em; width: 5px; height: 5px; border-radius: 50%; background: var(--resume-accent); }}
    @media (max-width: 640px) {{
      .resume-preview-shell {{ padding: 8px; }}
      .resume-preview-page {{ min-height: 0; padding: 28px 24px 38px; }}
      .resume-preview-title {{ margin: -28px -24px 18px; }}
      .resume-preview-header {{ padding: 18px; }}
      .resume-preview-name {{ font-size: 25px; }}
      .resume-preview-photo {{ width: 70px; height: 88px; }}
    }}
  </style>
  <article class="resume-preview-page" aria-label="{html.escape(template['name'])}简历实时预览">
    <div class="resume-preview-title">简历</div>
    <header class="{header_class}">
      <div>
        <h2 class="resume-preview-name">{name}</h2>
        <p class="resume-preview-role">目标岗位：{role}</p>
        {f'<p class="resume-preview-contact">{contact}</p>' if contact else ''}
      </div>
      {photo_html}
    </header>
    <div class="resume-preview-body">{''.join(body_parts)}</div>
  </article>
</div>
"""


def _set_word_font(run, size: float = 10.5, bold: bool = False, color: str = "202938") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def _word_text(paragraph, text: str, size: float = 10.5, bold: bool = False, color: str = "202938") -> None:
    _set_word_font(paragraph.add_run(text), size=size, bold=bold, color=color)


def _shade_word_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _remove_word_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_word_heading_style(paragraph, color: str) -> None:
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    properties.append(borders)


def create_resume_document(
    resume_data: dict,
    optimized_text: str | None = None,
    photo_bytes: bytes | None = None,
    template_id: str = DEFAULT_TEMPLATE_ID,
) -> bytes:
    """生成所选专业模板的完整 Word 简历。"""
    template = _template(template_id)
    primary = template["primary"].lstrip("#")
    accent_color = template["accent"].lstrip("#")
    pale = template["pale"].lstrip("#")
    text_color = template["text"].lstrip("#")
    muted = template["muted"].lstrip("#")
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    accent = document.add_table(rows=1, cols=1)
    accent.autofit = False
    accent.columns[0].width = Cm(17.0)
    _remove_word_table_borders(accent)
    _shade_word_cell(accent.cell(0, 0), primary)
    accent.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(0)
    _word_text(accent.cell(0, 0).paragraphs[0], " ", 2)

    header = document.add_table(rows=1, cols=2)
    header.autofit = False
    header.columns[0].width = Cm(13.7)
    header.columns[1].width = Cm(3.3)
    _remove_word_table_borders(header)
    left_cell, right_cell = header.rows[0].cells
    header_fill = primary if template["header_mode"] == "dark" else pale
    _shade_word_cell(left_cell, header_fill)
    _shade_word_cell(right_cell, header_fill)
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    name_paragraph = left_cell.paragraphs[0]
    header_text = "FFFFFF" if template["header_mode"] == "dark" else primary
    header_muted = "DCE7EF" if template["header_mode"] == "dark" else muted
    _word_text(name_paragraph, _plain(resume_data.get("name")) or "个人简历", 22, True, header_text)
    name_paragraph.paragraph_format.space_after = Pt(4)
    role_paragraph = left_cell.add_paragraph()
    role_color = "E8EEF4" if template["header_mode"] == "dark" else accent_color
    _word_text(role_paragraph, f"目标岗位：{_plain(resume_data.get('target_role'))}", 11, True, role_color)
    role_paragraph.paragraph_format.space_after = Pt(3)
    contact = _contact_line(resume_data)
    if contact:
        contact_paragraph = left_cell.add_paragraph()
        _word_text(contact_paragraph, contact, 9.5, False, header_muted)

    if photo_bytes:
        photo_paragraph = right_cell.paragraphs[0]
        photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        photo_paragraph.add_run().add_picture(BytesIO(photo_bytes), width=Cm(2.8))

    document.add_paragraph().paragraph_format.space_after = Pt(0)
    for block_type, text in _content_blocks(resume_data, optimized_text):
        if block_type == "heading":
            paragraph = document.add_paragraph()
            _set_word_heading_style(paragraph, accent_color)
            _word_text(paragraph, text, 13, True, primary)
        elif block_type == "bullet":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.35)
            paragraph.paragraph_format.first_line_indent = Cm(-0.25)
            paragraph.paragraph_format.space_after = Pt(2)
            _word_text(paragraph, f"· {text}", color=text_color)
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(3)
            _word_text(paragraph, text, color=text_color)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _word_text(footer, f"个人简历 ｜ {template['name']}", 8, False, muted)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _register_pdf_font() -> str:
    try:
        pdfmetrics.getFont(EMBEDDED_CHINESE_FONT)
        return EMBEDDED_CHINESE_FONT
    except KeyError:
        pass

    for font_path in FONT_CANDIDATES:
        if not font_path.is_file():
            continue
        try:
            pdfmetrics.registerFont(
                TTFont(EMBEDDED_CHINESE_FONT, str(font_path), subfontIndex=0)
            )
            return EMBEDDED_CHINESE_FONT
        except Exception:
            continue

    try:
        pdfmetrics.getFont(CHINESE_FONT)
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont(CHINESE_FONT))
    return CHINESE_FONT


def _register_pdf_fonts() -> tuple[str, str]:
    """注册屏幕和打印都清晰的中文常规/粗体字体。"""
    regular_font = _register_pdf_font()
    try:
        pdfmetrics.getFont(EMBEDDED_CHINESE_BOLD_FONT)
        return regular_font, EMBEDDED_CHINESE_BOLD_FONT
    except KeyError:
        pass

    for font_path in BOLD_FONT_CANDIDATES:
        if not font_path.is_file():
            continue
        try:
            pdfmetrics.registerFont(
                TTFont(
                    EMBEDDED_CHINESE_BOLD_FONT,
                    str(font_path),
                    subfontIndex=0,
                )
            )
            return regular_font, EMBEDDED_CHINESE_BOLD_FONT
        except Exception:
            continue
    return regular_font, regular_font


def _pdf_photo(photo_bytes: bytes) -> PdfImage:
    photo_buffer = BytesIO(photo_bytes)
    image_reader = ImageReader(photo_buffer)
    pixel_width, pixel_height = image_reader.getSize()
    max_width, max_height = 28 * mm, 38 * mm
    scale = min(max_width / pixel_width, max_height / pixel_height)
    photo_buffer.seek(0)
    return PdfImage(photo_buffer, width=pixel_width * scale, height=pixel_height * scale)


class SectionIcon(Flowable):
    """用矢量线条绘制章节图标，避免依赖可能缺字的图标字体。"""

    def __init__(self, kind: str, template: dict):
        super().__init__()
        self.kind = kind
        self.template = template
        self.width = 11 * mm
        self.height = 11 * mm

    def wrap(self, available_width: float, available_height: float):
        return self.width, self.height

    def draw(self) -> None:
        canvas = self.canv
        center = self.width / 2
        radius = 4.6 * mm
        canvas.saveState()
        canvas.setFillColor(colors.HexColor(self.template["soft"]))
        canvas.setStrokeColor(colors.HexColor(self.template["border"]))
        canvas.setLineWidth(0.7)
        canvas.circle(center, self.height / 2, radius, fill=1, stroke=1)
        canvas.setStrokeColor(colors.HexColor(self.template["accent"]))
        canvas.setFillColor(colors.HexColor(self.template["accent"]))
        canvas.setLineWidth(1.05)

        if self.kind == "education":
            canvas.line(center - 3.4 * mm, center + 0.8 * mm, center, center + 2.6 * mm)
            canvas.line(center, center + 2.6 * mm, center + 3.4 * mm, center + 0.8 * mm)
            canvas.line(center - 2.5 * mm, center + 0.1 * mm, center + 2.5 * mm, center + 0.1 * mm)
            canvas.line(center - 2.1 * mm, center - 0.3 * mm, center - 2.1 * mm, center - 2.1 * mm)
            canvas.line(center + 2.1 * mm, center - 0.3 * mm, center + 2.1 * mm, center - 2.1 * mm)
        elif self.kind == "work":
            canvas.roundRect(center - 3.1 * mm, center - 2.2 * mm, 6.2 * mm, 4.2 * mm, 0.7 * mm, fill=0, stroke=1)
            canvas.arc(center - 1.6 * mm, center + 0.8 * mm, center + 1.6 * mm, center + 3.2 * mm, 0, 180)
            canvas.line(center - 3.1 * mm, center - 0.2 * mm, center + 3.1 * mm, center - 0.2 * mm)
            canvas.line(center, center + 0.4 * mm, center, center - 0.8 * mm)
        elif self.kind == "project":
            points = [
                (center - 2.8 * mm, center + 2 * mm),
                (center + 2.8 * mm, center + 1.1 * mm),
                (center - 1.6 * mm, center - 2.2 * mm),
            ]
            canvas.line(*points[0], *points[1])
            canvas.line(*points[1], *points[2])
            canvas.line(*points[2], *points[0])
            for x, y in points:
                canvas.circle(x, y, 1.1 * mm, fill=1, stroke=0)
        elif self.kind == "skills":
            canvas.roundRect(center - 3.1 * mm, center - 3.1 * mm, 6.2 * mm, 6.2 * mm, 1 * mm, fill=0, stroke=1)
            canvas.line(center - 2 * mm, center, center - 0.5 * mm, center - 1.6 * mm)
            canvas.line(center - 0.5 * mm, center - 1.6 * mm, center + 2.3 * mm, center + 1.8 * mm)
        else:
            for offset in (1.7 * mm, 0, -1.7 * mm):
                canvas.line(center - 2.5 * mm, center + offset, center + 2.5 * mm, center + offset)
        canvas.restoreState()


def _pdf_section_heading(title: str, kind: str, style, width: float, template: dict) -> Table:
    heading = Table(
        [[SectionIcon(kind, template), Paragraph(html.escape(title), style)]],
        colWidths=[12 * mm, width - 12 * mm],
    )
    heading.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(template["pale"])),
        ("BOX", (0, 0), (-1, -1), 0.55, colors.HexColor(template["border"])),
        ("LINEBELOW", (0, 0), (-1, -1), 0.8, colors.HexColor(template["border"])),
        ("LEFTPADDING", (0, 0), (0, 0), 5),
        ("RIGHTPADDING", (0, 0), (0, 0), 1),
        ("LEFTPADDING", (1, 0), (1, 0), 3),
        ("RIGHTPADDING", (1, 0), (1, 0), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return heading


def _draw_pdf_page(canvas, doc, pdf_font: str, template: dict) -> None:
    """绘制所选模板的页面背景和页码。"""
    canvas.saveState()
    page_width, page_height = A4
    canvas.setFillColor(colors.white)
    canvas.rect(0, 0, page_width, page_height, fill=1, stroke=0)
    canvas.setFillColor(colors.HexColor(template["primary"]))
    canvas.rect(0, page_height - 4 * mm, page_width, 4 * mm, fill=1, stroke=0)
    canvas.setFillColor(colors.HexColor(template["soft"]))
    canvas.circle(page_width + 5 * mm, page_height - 18 * mm, 26 * mm, fill=1, stroke=0)
    canvas.setFillColor(colors.HexColor(template["pale"]))
    canvas.circle(-8 * mm, 17 * mm, 22 * mm, fill=1, stroke=0)
    canvas.setStrokeColor(colors.HexColor(template["border"]))
    canvas.setLineWidth(0.6)
    canvas.line(16 * mm, 12 * mm, page_width - 16 * mm, 12 * mm)
    canvas.setFont(pdf_font, 8)
    canvas.setFillColor(colors.HexColor(template["muted"]))
    canvas.drawString(17 * mm, 7.5 * mm, f"个人简历 ｜ {template['name']}")
    canvas.drawRightString(page_width - 17 * mm, 7.5 * mm, f"第 {doc.page} 页")
    canvas.restoreState()


def create_resume_pdf(
    resume_data: dict,
    optimized_text: str | None = None,
    photo_bytes: bytes | None = None,
    template_id: str = DEFAULT_TEMPLATE_ID,
) -> bytes:
    """生成所选专业模板、适合招聘阅读和打印的中文 PDF 简历。"""
    template = _template(template_id)
    pdf_primary = colors.HexColor(template["primary"])
    pdf_accent = colors.HexColor(template["accent"])
    pdf_pale = colors.HexColor(template["pale"])
    pdf_border = colors.HexColor(template["border"])
    pdf_text = colors.HexColor(template["text"])
    pdf_muted = colors.HexColor(template["muted"])
    pdf_font, pdf_bold_font = _register_pdf_fonts()
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=12 * mm,
        bottomMargin=16 * mm,
        title=f"{_plain(resume_data.get('name'))}的简历",
        author=_plain(resume_data.get("name")),
    )
    styles = getSampleStyleSheet()
    document_title_style = ParagraphStyle(
        "ResumeDocumentTitle",
        parent=styles["Title"],
        fontName=pdf_bold_font,
        fontSize=15,
        leading=19,
        alignment=TA_CENTER,
        textColor=colors.white,
        spaceAfter=0,
    )
    name_style = ParagraphStyle(
        "ResumeName", parent=styles["Title"], fontName=pdf_bold_font,
        fontSize=23, leading=28, alignment=TA_LEFT,
        textColor=colors.white if template["header_mode"] == "dark" else pdf_primary,
        spaceAfter=4,
    )
    role_style = ParagraphStyle(
        "ResumeRole", parent=styles["Normal"], fontName=pdf_bold_font,
        fontSize=10.8, leading=16, textColor=pdf_accent, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "ResumeContact", parent=styles["Normal"], fontName=pdf_font,
        fontSize=9.4, leading=14.5,
        textColor=colors.HexColor("#DCE7EF") if template["header_mode"] == "dark" else pdf_muted,
    )
    heading_style = ParagraphStyle(
        "ResumeHeading", parent=styles["Heading2"], fontName=pdf_bold_font,
        fontSize=12.5, leading=17, textColor=pdf_primary,
        spaceBefore=0, spaceAfter=0,
    )
    body_style = ParagraphStyle(
        "ResumeBody", parent=styles["BodyText"], fontName=pdf_font,
        fontSize=10.4, leading=16.2, alignment=TA_LEFT,
        textColor=pdf_text, spaceAfter=4.2, wordWrap="CJK",
        splitLongWords=False, allowWidows=0, allowOrphans=0,
    )
    entry_style = ParagraphStyle(
        "ResumeEntry", parent=body_style, fontName=pdf_bold_font,
        textColor=pdf_primary,
        spaceBefore=0.8, spaceAfter=3.6,
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet", parent=body_style,
        leftIndent=5 * mm, firstLineIndent=0,
        bulletIndent=1.2 * mm, bulletFontName=pdf_bold_font,
        bulletFontSize=6.5, bulletColor=pdf_accent,
        spaceAfter=3.6,
    )

    header_items = [
        Paragraph(html.escape(_plain(resume_data.get("name")) or "姓名"), name_style),
        Paragraph(html.escape(f"目标岗位：{_plain(resume_data.get('target_role'))}"), role_style),
    ]
    contact = _contact_line(resume_data)
    if contact:
        header_items.append(Paragraph(html.escape(contact), contact_style))

    title_bar = Table(
        [[Paragraph("简历", document_title_style)]],
        colWidths=[document.width],
    )
    title_bar.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), pdf_primary),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4.2 * mm),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4.2 * mm),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
    ]))

    story = [title_bar, Spacer(1, 2.5 * mm)]
    if photo_bytes:
        header = Table(
            [[header_items, _pdf_photo(photo_bytes)]],
            colWidths=[document.width - 34 * mm, 34 * mm],
        )
    else:
        header = Table([[header_items]], colWidths=[document.width])
    header_commands = [
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BACKGROUND", (0, 0), (-1, -1), pdf_primary if template["header_mode"] == "dark" else pdf_pale),
        ("LEFTPADDING", (0, 0), (-1, -1), 7 * mm),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7 * mm),
        ("TOPPADDING", (0, 0), (-1, -1), 5.5 * mm),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5.5 * mm),
        ("BOX", (0, 0), (-1, -1), 0.65, pdf_border),
        ("LINEBELOW", (0, 0), (-1, -1), 1.2, pdf_accent),
    ]
    if photo_bytes:
        header_commands.extend([
            ("ALIGN", (1, 0), (1, 0), "RIGHT"),
            ("LINEBEFORE", (1, 0), (1, 0), 0.5, pdf_border),
            ("LEFTPADDING", (1, 0), (1, 0), 2.5 * mm),
            ("RIGHTPADDING", (1, 0), (1, 0), 2.5 * mm),
            ("TOPPADDING", (1, 0), (1, 0), 2.5 * mm),
            ("BOTTOMPADDING", (1, 0), (1, 0), 2.5 * mm),
        ])
    header.setStyle(TableStyle(header_commands))
    story.extend([header, Spacer(1, 3.2 * mm)])

    section_kinds = {
        "个人简介": "summary",
        "教育背景": "education",
        "工作经历": "work",
        "项目经历": "project",
        "技能证书": "skills",
    }
    current_heading: Table | None = None
    current_section_flowables: list[Flowable] = []
    section_groups: list[tuple[Table, list[Flowable]]] = []
    for block_type, text in _content_blocks(resume_data, optimized_text):
        safe_text = html.escape(text)
        if block_type == "heading":
            if current_heading is not None:
                section_groups.append((current_heading, current_section_flowables))
            current_heading = _pdf_section_heading(
                text,
                section_kinds.get(text, "summary"),
                heading_style,
                document.width,
                template,
            )
            current_section_flowables = []
        elif block_type == "bullet":
            content_flowable = Paragraph(safe_text, bullet_style, bulletText="●")
            if current_heading is not None:
                current_section_flowables.append(content_flowable)
            else:
                story.append(content_flowable)
        else:
            is_entry_line = bool(
                re.match(r"^(?:19|20)\d{2}(?:[./-]|年)", text)
                or ("｜" in text and len(text) <= 60)
            )
            content_flowable = Paragraph(
                safe_text,
                entry_style if is_entry_line else body_style,
            )
            if current_heading is not None:
                current_section_flowables.append(content_flowable)
            else:
                story.append(content_flowable)
    if current_heading is not None:
        section_groups.append((current_heading, current_section_flowables))

    for section_heading, section_flowables in section_groups:
        section_prefix: list[Flowable] = [
            Spacer(1, 2.7 * mm),
            section_heading,
            Spacer(1, 1.5 * mm),
        ]
        if not section_flowables:
            story.extend(section_prefix)
        elif len(section_flowables) <= 7:
            story.append(KeepTogether(section_prefix + section_flowables))
        else:
            story.append(KeepTogether(section_prefix + [section_flowables[0]]))
            story.extend(section_flowables[1:])

    document.build(
        story,
        onFirstPage=lambda canvas, doc: _draw_pdf_page(canvas, doc, pdf_font, template),
        onLaterPages=lambda canvas, doc: _draw_pdf_page(canvas, doc, pdf_font, template),
    )
    return output.getvalue()
